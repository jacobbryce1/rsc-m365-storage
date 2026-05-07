#!/usr/bin/env python3
"""
Generate documentation:
1. Word document (wiki-style deployment & operations guide)
2. Updated README.md

Changes v2.1:
- Project structure tree updated: SECURITY.md, generate_docs.py, package_project.py added
- Security section updated to reflect v2.1 hardening (10-item findings table)
- TLS configuration section updated (RSC_CA_BUNDLE, verify=False removed)
- .env instructions include chmod 600 guidance
- Full Configuration Reference table added to Word doc
- Troubleshooting table updated with v2.1 error messages
"""

import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    print("Installing python-docx...")
    subprocess.run(["pip", "install", "python-docx"], check=True)
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_code_block(doc, code):
    """Add a formatted code block."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    para.paragraph_format.element.get_or_add_pPr().append(shd)


def generate_word_doc():
    """Generate the Word document."""
    doc = Document()

    # Title
    title = doc.add_heading("RSC M365 Licensing Compliance & Storage Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Deployment, Operations & Interpretation Guide")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Version 2.1 - " + datetime.now().strftime("%B %d, %Y")).font.size = Pt(11)

    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. Overview & Purpose",
        "2. M365 Licensing Model in RSC",
        "3. Report Outputs & Interpretation",
        "4. Prerequisites",
        "5. Deployment - macOS",
        "6. Deployment - Linux",
        "7. Deployment - Windows",
        "8. Configuration",
        "9. Full Configuration Reference",
        "10. Running the Report",
        "11. Scheduling (Automated Execution)",
        "12. Data Sources & API Details",
        "13. Security",
        "14. Troubleshooting",
        "15. Appendix: RSC GraphQL Queries Used",
    ]
    for item in toc:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 1. Overview
    # -------------------------------------------------------------------------
    doc.add_heading("1. Overview & Purpose", level=1)
    doc.add_paragraph(
        "This tool retrieves Microsoft 365 backup licensing and storage consumption data "
        "from Rubrik Security Cloud (RSC) via the GraphQL API. It produces a comprehensive "
        "report that helps organizations:"
    )
    for b in [
        "Monitor license compliance (entitled vs consumed users and capacity)",
        "Track monthly growth in protected users and data volumes",
        "Break down consumption by M365 workload type (Exchange, OneDrive, SharePoint, Teams)",
        "Forecast future capacity needs based on historical trends",
        "Identify potential license overages before they occur",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Key Metrics Provided", level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["Metric", "Granularity", "History", "Use Case"]):
        table.rows[0].cells[i].text = h
    metrics = [
        ("License Compliance", "Account-wide", "Current", "Are we within entitlement?"),
        ("Consumed User Licenses", "Monthly x Type", "~12 months", "User license trending"),
        ("Data Protection Capacity", "Monthly x Type", "~12 months", "DPC/capacity growth"),
        ("Physical Storage", "Per Org", "10 days", "Backend storage efficiency"),
        ("Growth Projections", "Calculated", "Forward 12mo", "Capacity planning"),
    ]
    for i, (m, g, h, u) in enumerate(metrics, 1):
        table.rows[i].cells[0].text = m
        table.rows[i].cells[1].text = g
        table.rows[i].cells[2].text = h
        table.rows[i].cells[3].text = u
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 2. Licensing Model
    # -------------------------------------------------------------------------
    doc.add_heading("2. M365 Licensing Model in RSC", level=1)
    doc.add_paragraph(
        "Rubrik Security Cloud measures M365 protection licensing on two primary metrics:"
    )

    doc.add_heading("User Licenses (per-user, per-month)", level=2)
    doc.add_paragraph(
        "Priced as $ per user per month. Editions include M365 Foundation, Professional, "
        "and Enterprise (plus EDU variants), each with different feature sets."
    )
    doc.add_paragraph("How consumed users are calculated:")
    doc.add_paragraph(
        "Consumed Users = MAX(protected user mailboxes, protected shared mailboxes, protected OneDrives)",
        style="List Bullet",
    )
    doc.add_paragraph(
        "This report tracks mailbox and OneDrive counts monthly and computes the consumed "
        "user count using this formula."
    )

    doc.add_heading("Data Protection Capacity (DPC, per-GB per-month)", level=2)
    doc.add_paragraph(
        "Separate capacity SKUs priced as $ per GB per month. RSC tracks DPC as the total "
        "front-end GB of live M365 snapshot data protected across Exchange, OneDrive, "
        "SharePoint, and Teams."
    )
    doc.add_paragraph(
        "This report uses 'Data Transferred' (transferredBytes) as the DPC metric — "
        "this represents the cumulative front-end data ingested during backup operations "
        "per workload type per month, which directly corresponds to the protected data volume."
    )

    doc.add_heading("Compliance Status", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Status"
    table.rows[0].cells[1].text = "Meaning"
    table.rows[1].cells[0].text = "OK"
    table.rows[1].cells[1].text = "Utilization <= 100% — within entitlement"
    table.rows[2].cells[0].text = "OVER"
    table.rows[2].cells[1].text = "Utilization > 100% — exceeds entitlement, action needed"
    table.rows[3].cells[0].text = "Unknown"
    table.rows[3].cells[1].text = "Entitlement data not available (check service account permissions)"
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 3. Report Outputs
    # -------------------------------------------------------------------------
    doc.add_heading("3. Report Outputs & Interpretation", level=1)
    doc.add_heading("Console Output Sections", level=2)
    sections = [
        ("Section 1: License Compliance",
         "Shows entitled vs consumed for both User Licenses and DPC at the account level, "
         "with utilization percentages and compliance status. Also breaks down consumption per org."),
        ("Section 2: Monthly Consumed User Licenses",
         "12-month table showing monthly counts of protected Mailboxes, OneDrives, SharePoint Sites, "
         "and Teams. The 'Consumed Users' column shows max(Mailboxes, OneDrives) which is the "
         "billable user count per Rubrik's licensing formula."),
        ("Section 3: Monthly DPC by Workload Type",
         "12-month table showing data transferred (in GB) per workload type. Columns: Exchange, "
         "OneDrive, SharePoint, Teams, and Total. This represents the monthly data protection "
         "capacity consumption."),
        ("Section 4: Growth Analysis",
         "Calculates user growth rate, DPC growth rate, average monthly growth, and projected "
         "annual capacity. Use this for license renewal planning."),
        ("Section 5: Physical Storage Trend",
         "10-day trend of actual backend physical storage consumed across all M365 orgs. "
         "Shows the post-dedup/compression storage footprint."),
    ]
    for sect_title, desc in sections:
        doc.add_heading(sect_title, level=3)
        doc.add_paragraph(desc)

    doc.add_heading("Excel Export Sheets", level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Sheet Name"
    table.rows[0].cells[1].text = "Contents"
    sheets = [
        ("License Compliance", "Entitled vs consumed summary with utilization % and status"),
        ("Monthly Users", "Monthly protected object counts and consumed user calculation"),
        ("Monthly DPC (GB)", "Monthly data transferred per workload type in GB"),
        ("Monthly DPC (Bytes)", "Same data in raw bytes for precision calculations"),
        ("Per-Org Consumption", "Current consumption broken down by M365 organization"),
        ("10-Day Trend", "Daily physical storage values for the last 10 days"),
    ]
    for i, (name, desc) in enumerate(sheets, 1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = desc
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 4. Prerequisites
    # -------------------------------------------------------------------------
    doc.add_heading("4. Prerequisites", level=1)
    doc.add_heading("Software", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Component", "Version", "Notes"]):
        table.rows[0].cells[i].text = h
    prereqs = [
        ("Python", "3.9+", "3.12 recommended"),
        ("pip", "Latest", "Included with Python"),
        ("Network", "HTTPS 443", "Outbound to *.my.rubrik.com"),
    ]
    for i, (c, v, n) in enumerate(prereqs, 1):
        table.rows[i].cells[0].text = c
        table.rows[i].cells[1].text = v
        table.rows[i].cells[2].text = n

    doc.add_heading("RSC Service Account Requirements", level=2)
    doc.add_paragraph("Create a service account in RSC with the following permissions:")
    for p in [
        "Read access to M365 organizations and objects",
        "Read access to Reports and Compliance data",
        "View license entitlement information",
    ]:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("Creating the Service Account", level=3)
    for i, s in enumerate([
        "Log into RSC at https://YOUR_ORG.my.rubrik.com",
        "Navigate to Settings → Service Accounts",
        "Click 'Create Service Account'",
        "Name: 'm365-reporting' (or similar descriptive name)",
        "Assign role: read-only access to M365 objects and Reports",
        "Copy the Client ID and Client Secret immediately (shown only once)",
        "Store credentials in .env — never commit to source control",
    ], 1):
        doc.add_paragraph(f"{i}. {s}")
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 5. macOS Deployment
    # -------------------------------------------------------------------------
    doc.add_heading("5. Deployment - macOS", level=1)
    doc.add_heading("Quick Start", level=2)
    add_code_block(doc,
        "unzip rsc-m365-storage-report-YYYYMMDD.zip\n"
        "cd rsc-m365-storage-report\n\n"
        "chmod +x deploy.sh\n"
        "./deploy.sh\n\n"
        "# Configure credentials\n"
        "nano .env\n\n"
        "# Run the report\n"
        "python run_report.py"
    )
    doc.add_heading("What deploy.sh Does", level=3)
    for s in [
        "Verifies Python 3.9+ is available",
        "Creates a Python virtual environment (venv/)",
        "Installs all dependencies from requirements.txt",
        "Copies .env.example to .env and sets permissions to 600 (owner-only)",
    ]:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 6. Linux Deployment
    # -------------------------------------------------------------------------
    doc.add_heading("6. Deployment - Linux", level=1)
    doc.add_heading("Supported Distributions", level=2)
    for d in ["Ubuntu 20.04+", "Debian 11+", "RHEL / Rocky / Alma 8+", "Fedora 36+", "CentOS Stream 8+"]:
        doc.add_paragraph(d, style="List Bullet")
    doc.add_heading("Quick Start", level=2)
    add_code_block(doc,
        "unzip rsc-m365-storage-report-YYYYMMDD.zip\n"
        "cd rsc-m365-storage-report\n"
        "chmod +x deploy_linux.sh\n"
        "./deploy_linux.sh"
    )
    doc.add_paragraph(
        "The Linux deployment script auto-detects your distribution and installs "
        "python3-venv and python3-pip if needed (requires sudo for package installation). "
        "It also optionally configures a monthly cron job."
    )
    doc.add_heading("Monthly Cron Job", level=2)
    add_code_block(doc,
        "crontab -e\n"
        "# Add this line (runs 1st of each month at 6 AM):\n"
        "0 6 1 * * cd /path/to/rsc-m365-storage-report && "
        "./venv/bin/python final_m365_report.py >> output/cron.log 2>&1"
    )
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 7. Windows Deployment
    # -------------------------------------------------------------------------
    doc.add_heading("7. Deployment - Windows", level=1)
    doc.add_heading("Step 1: Install Python", level=2)
    doc.add_paragraph(
        "Download Python 3.9+ from https://www.python.org/downloads/windows/. "
        "IMPORTANT: During installation, check the box 'Add Python to PATH'."
    )
    doc.add_heading("Step 2: Extract & Deploy", level=2)
    add_code_block(doc,
        "REM Extract zip via Windows Explorer (right-click → Extract All)\n"
        "REM Open Command Prompt in the extracted folder:\n"
        "cd rsc-m365-storage-report\n"
        "deploy.bat"
    )
    doc.add_heading("Step 3: Configure", level=2)
    add_code_block(doc, "notepad .env\n# Set RSC_URL, RSC_CLIENT_ID, RSC_CLIENT_SECRET")
    doc.add_heading("Step 4: Run", level=2)
    add_code_block(doc, "python run_report.py")
    doc.add_heading("Task Scheduler (Monthly Automation)", level=2)
    for s in [
        "Open Task Scheduler (Win+R → taskschd.msc)",
        "Create Basic Task → Name: 'RSC M365 License Report'",
        "Trigger: Monthly, Day 1, 6:00 AM",
        "Action: Start a Program",
        "Program: C:\\path\\to\\venv\\Scripts\\python.exe",
        "Arguments: final_m365_report.py",
        "Start in: C:\\path\\to\\rsc-m365-storage-report",
    ]:
        doc.add_paragraph(s, style="List Number")
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 8. Configuration
    # -------------------------------------------------------------------------
    doc.add_heading("8. Configuration", level=1)
    doc.add_heading("Creating Your .env File", level=2)
    add_code_block(doc, "cp .env.example .env\nchmod 600 .env   # macOS/Linux only")
    doc.add_paragraph(
        "Never commit .env to version control. It is listed in .gitignore. "
        "The deploy scripts set file permissions to 600 (owner-read/write only) automatically on macOS and Linux."
    )

    doc.add_heading("Environment Variables (.env file)", level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["Variable", "Required", "Format", "Description"]):
        table.rows[0].cells[i].text = h
    for i, (v, r, f, d) in enumerate([
        ("RSC_URL",         "Yes",      "https://org.my.rubrik.com", "RSC instance URL — must match *.my.rubrik.com"),
        ("RSC_CLIENT_ID",   "Yes",      "UUID",                     "Service account client ID"),
        ("RSC_CLIENT_SECRET","Yes",     "String",                   "Service account secret"),
        ("RSC_CA_BUNDLE",   "No",       "true | /path/to/file.pem", "CA bundle override for corporate PKI"),
    ], 1):
        table.rows[i].cells[0].text = v
        table.rows[i].cells[1].text = r
        table.rows[i].cells[2].text = f
        table.rows[i].cells[3].text = d
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 9. Full Configuration Reference
    # -------------------------------------------------------------------------
    doc.add_heading("9. Full Configuration Reference", level=1)

    doc.add_heading("Required Variables", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["Variable", "Description"]):
        table.rows[0].cells[i].text = h
    for i, (v, d) in enumerate([
        ("RSC_URL",          "Base URL of your RSC tenant — must match https://<org>.my.rubrik.com"),
        ("RSC_CLIENT_ID",    "RSC service account client ID"),
        ("RSC_CLIENT_SECRET","RSC service account secret"),
    ], 1):
        table.rows[i].cells[0].text = v
        table.rows[i].cells[1].text = d

    doc.add_heading("Optional Variables", level=2)
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Variable", "Default", "Description"]):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "RSC_CA_BUNDLE"
    table.rows[1].cells[1].text = "true"
    table.rows[1].cells[2].text = "true (system CAs) or path to a .pem CA bundle for corporate PKI"
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 10. Running the Report
    # -------------------------------------------------------------------------
    doc.add_heading("10. Running the Report", level=1)
    doc.add_heading("Standard Execution", level=2)
    add_code_block(doc, "python run_report.py")
    doc.add_heading("Expected Runtime", level=2)
    doc.add_paragraph("30–60 seconds depending on network latency to RSC and number of orgs.")
    doc.add_heading("Output Location", level=2)
    add_code_block(doc, "output/m365_license_compliance_report_20260505_143022.xlsx")
    doc.add_heading("Manual Execution", level=2)
    add_code_block(doc,
        "# Activate virtual environment\n"
        "source venv/bin/activate        # macOS/Linux\n"
        "venv\\Scripts\\activate.bat      # Windows\n\n"
        "# Run directly\n"
        "python final_m365_report.py"
    )
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 11. Scheduling
    # -------------------------------------------------------------------------
    doc.add_heading("11. Scheduling (Automated Execution)", level=1)
    doc.add_paragraph(
        "For continuous compliance monitoring, schedule monthly execution. "
        "Each run produces a timestamped report file, building a historical archive."
    )
    doc.add_heading("macOS (launchd)", level=2)
    add_code_block(doc,
        "# Create ~/Library/LaunchAgents/com.rsc.m365report.plist\n"
        "# Key settings:\n"
        "#   ProgramArguments: /path/to/venv/bin/python final_m365_report.py\n"
        "#   WorkingDirectory: /path/to/rsc-m365-storage-report\n"
        "#   StartCalendarInterval: Day=1, Hour=6\n\n"
        "launchctl load ~/Library/LaunchAgents/com.rsc.m365report.plist"
    )
    doc.add_heading("Linux (cron)", level=2)
    add_code_block(doc,
        "0 6 1 * * cd /opt/rsc-m365-storage-report && "
        "./venv/bin/python final_m365_report.py >> output/cron.log 2>&1"
    )
    doc.add_heading("Windows", level=2)
    doc.add_paragraph("See Section 7 for Task Scheduler setup.")
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 12. Data Sources & API Details
    # -------------------------------------------------------------------------
    doc.add_heading("12. Data Sources & API Details", level=1)
    doc.add_heading("RSC GraphQL Queries Used", level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Query", "Returns", "Licensing Metric"]):
        table.rows[0].cells[i].text = h
    for i, (q, r, m) in enumerate([
        ("m365LicenseEntitlement",    "usersEntitled, capacityEntitledInBytes", "Entitled (what you pay for)"),
        ("o365Consumption",           "usersProtected, fetbConsumed per org",   "Current consumed (point-in-time)"),
        ("snappableGroupByConnection","Monthly counts + transferredBytes",       "Historical consumed (12 months)"),
        ("o365StorageStats",          "Physical/logical, 10-day trend",          "Backend storage efficiency"),
        ("o365Orgs",                  "M365 organization list",                  "Org enumeration"),
    ], 1):
        table.rows[i].cells[0].text = q
        table.rows[i].cells[1].text = r
        table.rows[i].cells[2].text = m

    doc.add_heading("Consumed User License Calculation", level=2)
    add_code_block(doc,
        "Consumed Users = MAX(O365Mailbox count, O365Onedrive count)\n\n"
        "Source: snappableGroupByConnection grouped by Month -> ObjectType\n"
        "  O365Mailbox count  = protected mailboxes\n"
        "  O365Onedrive count = protected OneDrives"
    )
    doc.add_heading("DPC Calculation", level=2)
    add_code_block(doc,
        "DPC = transferredBytes (cumulative data ingested per month)\n\n"
        "Source: snappableGroupByConnection -> aggregation.transferredBytes\n"
        "  Broken down by: Exchange, OneDrive, SharePoint, Teams\n"
        "  Reported in GB per month"
    )
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 13. Security (v2.1)
    # -------------------------------------------------------------------------
    doc.add_heading("13. Security", level=1)
    doc.add_paragraph(
        "v2.1 was reviewed against OWASP Top 10 (2021), NIST CSF 2.0, CIS Controls v8, "
        "and MITRE ATT&CK for Enterprise. The following hardening measures are in place."
    )

    for sec_title, desc in [
        ("Credential Protection",
         "Credentials are loaded exclusively from .env — never hard-coded. "
         "Tokens are kept in memory only and never written to disk or logs. "
         "Exception messages are sanitized; CLIENT_ID, CLIENT_SECRET, and token values "
         "never appear in tracebacks. Use a dedicated read-only service account and rotate "
         "the secret periodically."),
        ("TLS Verification",
         "All RSC API calls verify TLS against system CAs. verify=False has been removed "
         "from the codebase entirely. RSC_URL is validated at startup against "
         "https://*.my.rubrik.com — plaintext URLs and unexpected hosts are rejected before "
         "any credentials are transmitted. Corporate CA bundles can be specified via "
         "RSC_CA_BUNDLE without weakening verification."),
        ("Output File Security",
         "Output files contain M365 organization names, user counts, and data volumes. "
         "Treat as Internal / Confidential. The output/ directory is excluded from version "
         "control via .gitignore. Output files contain no credentials or tokens."),
        ("Dependency Auditing",
         "Dependencies are pinned in requirements.txt including urllib3 and certifi for "
         "TLS-layer patch tracking. Run: pip install --upgrade -r requirements.txt to pull "
         "CVE patches."),
    ]:
        doc.add_heading(sec_title, level=2)
        doc.add_paragraph(desc)

    doc.add_heading("Security Review Findings (v2.1)", level=2)
    table = doc.add_table(rows=11, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["#", "Framework", "Finding", "Status"]):
        table.rows[0].cells[i].text = h
    findings = [
        ("1",  "OWASP A02", "requests.post() lacked explicit verify= — silent TLS bypass possible",                     "Fixed"),
        ("2",  "OWASP A07", "os.getenv() returned None silently — unauthenticated API calls possible",                  "Fixed"),
        ("3",  "OWASP A05", "RSC_URL accepted any scheme/host — credentials could be sent to wrong host",               "Fixed"),
        ("4",  "OWASP A09", "No raise_for_status() — HTTP 401/403 silently returned None data",                         "Fixed"),
        ("5",  "NIST PR.DS-1", "Exception tracebacks could expose credential fragments",                                "Fixed"),
        ("6",  "OWASP A03", "Nested data['key'] access crashed on unexpected API response shapes",                      "Fixed"),
        ("7",  "NIST PR.DS-5", "output/*.docx, *.pdf, *.log not excluded from .gitignore",                             "Fixed"),
        ("8",  "CIS Control 4", "urllib3 and certifi absent from requirements.txt; versions stale",                     "Fixed"),
        ("9",  "NIST PR.DS-2", "os.makedirs missing exist_ok=True — crash if output/ absent",                          "Fixed"),
        ("10", "CIS Control 16", "os.path.join not used — forward-slash path non-portable on Windows",                 "Fixed"),
    ]
    for i, (n, fw, finding, status) in enumerate(findings, 1):
        table.rows[i].cells[0].text = n
        table.rows[i].cells[1].text = fw
        table.rows[i].cells[2].text = finding
        table.rows[i].cells[3].text = status
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 14. Troubleshooting
    # -------------------------------------------------------------------------
    doc.add_heading("14. Troubleshooting", level=1)
    for issue, fix in [
        ("RSC_URL is not set",
         "Copy .env.example to .env and set all three required variables."),
        ("RSC_URL must use HTTPS",
         "Verify the URL in .env starts with https://"),
        ("RSC_URL does not match expected pattern",
         "URL must be exactly https://<org>.my.rubrik.com"),
        ("TLS certificate verification failed",
         "Set RSC_CA_BUNDLE=/path/to/ca-bundle.crt for corporate CAs."),
        ("Authentication failed (HTTP 401)",
         "Check RSC_CLIENT_ID and RSC_CLIENT_SECRET in .env."),
        ("Authentication failed (HTTP 403)",
         "Service account needs read access to M365 objects and Reports in RSC."),
        ("Connection Refused / Timeout",
         "Verify RSC_URL in .env. Check VPN/network. Ensure HTTPS 443 is open outbound."),
        ("Storage shows 0",
         "Normal for per-object M365. RSC tracks M365 at org aggregate level. Use DPC/transferredBytes."),
        ("LibreSSL Warning (macOS)",
         "Cosmetic only. To suppress: brew install python@3.12 and recreate venv."),
        ("No Monthly Data",
         "Verify M365 backups are actively running. Data only appears from backup start date."),
        ("Python Not Found (Windows)",
         "Re-install Python with 'Add to PATH' checked. Or use full path to python.exe."),
    ]:
        doc.add_heading(issue, level=2)
        doc.add_paragraph(fix)
    doc.add_page_break()

    # -------------------------------------------------------------------------
    # 15. Appendix
    # -------------------------------------------------------------------------
    doc.add_heading("15. Appendix: RSC GraphQL Queries Used", level=1)
    doc.add_heading("License Entitlement", level=2)
    add_code_block(doc, "{\n  m365LicenseEntitlement {\n    usersEntitled\n    capacityEntitledInBytes\n  }\n}")

    doc.add_heading("Per-Org Consumption", level=2)
    add_code_block(doc,
        'query($input: O365ConsumptionInput!) {\n'
        '  o365Consumption(input: $input) {\n'
        '    consumption {\n'
        '      usersProtected\n'
        '      fetbConsumed\n'
        '    }\n'
        '    consumptionPerWorkloadType {\n'
        '      workloadType\n'
        '      consumption { usersProtected fetbConsumed }\n'
        '    }\n'
        '  }\n'
        '}\n'
        '# Variable: {"input": {"o365OrgId": "<org-uuid>"}}'
    )

    doc.add_heading("Monthly Licensing Data (Core Query)", level=2)
    add_code_block(doc,
        '{\n'
        '  snappableGroupByConnection(\n'
        '    first: 24\n'
        '    groupBy: Month\n'
        '    filter: {\n'
        '      objectType: [O365Mailbox, O365Onedrive, O365Site, O365Teams]\n'
        '    }\n'
        '    requestedAggregations: [TRANSFERRED_BYTES, Count]\n'
        '  ) {\n'
        '    edges {\n'
        '      node {\n'
        '        groupByInfo {\n'
        '          ... on TimeRangeWithUnit { start end }\n'
        '        }\n'
        '        snappableConnection {\n'
        '          count\n'
        '          aggregation { transferredBytes }\n'
        '        }\n'
        '        snappableGroupBy(groupBy: ObjectType) {\n'
        '          groupByInfo {\n'
        '            ... on ObjectType { enumValue }\n'
        '          }\n'
        '          snappableConnection {\n'
        '            count\n'
        '            aggregation { transferredBytes }\n'
        '          }\n'
        '        }\n'
        '      }\n'
        '    }\n'
        '  }\n'
        '}'
    )

    doc.add_heading("Storage Statistics", level=2)
    add_code_block(doc,
        'query($orgID: UUID) {\n'
        '  o365StorageStats(orgID: $orgID) {\n'
        '    liveDataSizeInBytes\n'
        '    physicalDataSizeInBytes\n'
        '    storageEfficiencyPercent\n'
        '    dailyGrowthInBytes\n'
        '    estimatedThirtyDaysStorageInBytes\n'
        '    physicalDataSizeTimeSeries {\n'
        '      physicalDataSizeInBytes\n'
        '      timestamp\n'
        '    }\n'
        '  }\n'
        '}'
    )

    # Save
    os.makedirs("output", exist_ok=True)
    doc_path = os.path.join("output", "RSC_M365_License_Compliance_Report_Guide.docx")
    doc.save(doc_path)
    return doc_path


def generate_readme():
    """
    Generate README.md.
    NOTE: This regenerates the README from a template. If you have made manual
    edits to README.md (e.g. the v2.1 security review table), review the output
    before committing. The generated README matches the v2.1 structure.
    """
    lines = []

    lines.append("# \U0001f4ca RSC M365 Licensing Compliance & Storage Report")
    lines.append("")
    lines.append("Automated Microsoft 365 backup licensing compliance and storage consumption reporting via the Rubrik Security Cloud (RSC) GraphQL API. Designed for environments from single-org to multi-org M365 tenants managed through RSC.")
    lines.append("")
    lines.append("> **Not affiliated with Rubrik.** This is an independent, community-built tool. See [Legal & Disclaimer](#legal--disclaimer) for details.")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Overview")
    lines.append("")
    lines.append("This tool connects to your RSC instance, retrieves M365 backup licensing and storage data, and produces a comprehensive compliance and trending report — surfacing license overages, consumption trends, and growth projections as a timestamped Excel workbook with console summary output.")
    lines.append("")
    lines.append("**v2.1** introduces a full security hardening pass reviewed against OWASP Top 10, NIST CSF 2.0, CIS Controls v8, and MITRE ATT&CK. See [Security](#security) for details.")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Features")
    lines.append("")
    lines.append("| Feature | Details |")
    lines.append("|---------|---------|")
    lines.append("| \U0001f4cb **License compliance** | Entitled vs consumed for User Licenses and Data Protection Capacity (DPC) |")
    lines.append("| \U0001f4c5 **Monthly user trending** | Protected mailboxes, OneDrives, SharePoint sites, and Teams over 12 months |")
    lines.append("| \U0001f4be **DPC by workload type** | Data transferred per workload (Exchange, OneDrive, SharePoint, Teams) |")
    lines.append("| \U0001f4c8 **Growth analysis** | User growth rate, DPC growth rate, monthly average, annual projection |")
    lines.append("| \U0001f5c4\ufe0f **Physical storage trend** | 10-day backend storage with dedup efficiency |")
    lines.append("| \U0001f3e2 **Multi-org support** | Aggregates consumption across all M365 organizations in your RSC tenant |")
    lines.append("| \U0001f4e5 **Excel export** | Six-sheet timestamped `.xlsx` workbook saved to `output/` |")
    lines.append("| \U0001f512 **TLS-verified API calls** | Configurable CA bundle; `verify=False` removed entirely |")
    lines.append("| \U0001f511 **URL + credential validation** | RSC_URL validated against `*.my.rubrik.com` at startup; missing vars exit cleanly |")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## M365 Licensing Metrics")
    lines.append("")
    lines.append("| Metric | Formula | API Source |")
    lines.append("|--------|---------|------------|")
    lines.append("| Consumed Users | `max(Protected Mailboxes, Protected OneDrives)` | `snappableGroupByConnection` monthly counts |")
    lines.append("| Data Protection Capacity | Front-end data ingested per month | `transferredBytes` aggregation |")
    lines.append("| Entitled Users | Licensed user count | `m365LicenseEntitlement` |")
    lines.append("| Entitled Capacity | Licensed GB | `m365LicenseEntitlement` |")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Prerequisites")
    lines.append("")
    lines.append("| Requirement | Details |")
    lines.append("|-------------|---------|")
    lines.append("| Python | 3.9 or higher (3.12 recommended) |")
    lines.append("| Network | HTTPS access to your RSC instance (port 443 outbound) |")
    lines.append("| RSC Permissions | Service account with read access to M365 objects and Reports |")
    lines.append("| Disk Space | ~50 MB |")
    lines.append("")
    lines.append("> You must have a valid API key and an active Rubrik Security Cloud subscription. This tool does not bypass licensing or provide unauthorised access to any Rubrik features.")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Quick Start")
    lines.append("")
    lines.append("### macOS / Linux")
    lines.append("")
    lines.append("```bash")
    lines.append("unzip rsc-m365-storage-report-*.zip")
    lines.append("cd rsc-m365-storage-report")
    lines.append("chmod +x deploy.sh && ./deploy.sh")
    lines.append("cp .env.example .env")
    lines.append("nano .env   # Add your RSC credentials")
    lines.append("python run_report.py")
    lines.append("```")
    lines.append("")
    lines.append("### Windows (Command Prompt)")
    lines.append("")
    lines.append("```bat")
    lines.append("REM Extract zip, open Command Prompt in folder")
    lines.append("deploy.bat")
    lines.append("copy .env.example .env")
    lines.append("notepad .env   REM Add your RSC credentials")
    lines.append("python run_report.py")
    lines.append("```")
    lines.append("")
    lines.append("Reports are written to `output/` as timestamped Excel files when the run completes.")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Configuration")
    lines.append("")
    lines.append("### 1. Create your `.env` file")
    lines.append("")
    lines.append("```bash")
    lines.append("cp .env.example .env")
    lines.append("```")
    lines.append("")
    lines.append("```dotenv")
    lines.append("# Required")
    lines.append("RSC_URL=https://your-org.my.rubrik.com")
    lines.append("RSC_CLIENT_ID=your-client-id")
    lines.append("RSC_CLIENT_SECRET=your-client-secret")
    lines.append("")
    lines.append("# Optional: custom CA bundle for corporate proxies / private PKI")
    lines.append("# RSC_CA_BUNDLE=/path/to/ca-bundle.crt")
    lines.append("```")
    lines.append("")
    lines.append("> \u26a0\ufe0f **Never commit `.env` to version control.** It is already listed in `.gitignore`. On macOS/Linux, restrict the file to your user only: `chmod 600 .env`")
    lines.append("")
    lines.append("### 2. RSC Service Account Setup")
    lines.append("")
    lines.append("1. Log into RSC \u2192 **Settings** \u2192 **Service Accounts**")
    lines.append("2. Create a new service account")
    lines.append("3. Assign **read-only** access to M365 objects and Reports *(principle of least privilege)*")
    lines.append("4. Copy the Client ID and Secret \u2014 the secret is shown only once")
    lines.append("5. Paste values into your `.env`")
    lines.append("")
    lines.append("### 3. TLS Configuration")
    lines.append("")
    lines.append("RSC API calls verify TLS against system CAs by default. Override with `RSC_CA_BUNDLE`:")
    lines.append("")
    lines.append("```dotenv")
    lines.append("RSC_CA_BUNDLE=true                  # Default \u2014 verify against system CAs")
    lines.append("RSC_CA_BUNDLE=/path/to/bundle.pem   # Custom CA bundle for corporate PKI")
    lines.append("```")
    lines.append("")
    lines.append("> `verify=False` has been removed from this codebase entirely. TLS verification cannot be disabled.")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Full Configuration Reference")
    lines.append("")
    lines.append("### Required")
    lines.append("")
    lines.append("| Variable | Description |")
    lines.append("|----------|-------------|")
    lines.append("| `RSC_URL` | Base URL of your RSC tenant \u2014 must match `https://<org>.my.rubrik.com` |")
    lines.append("| `RSC_CLIENT_ID` | RSC service account client ID |")
    lines.append("| `RSC_CLIENT_SECRET` | RSC service account secret |")
    lines.append("")
    lines.append("### Optional")
    lines.append("")
    lines.append("| Variable | Default | Description |")
    lines.append("|----------|---------|-------------|")
    lines.append("| `RSC_CA_BUNDLE` | `true` | `true` (system CAs) or path to a `.pem` CA bundle |")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Output")
    lines.append("")
    lines.append("### Excel Workbook")
    lines.append("")
    lines.append("```")
    lines.append("output/m365_license_compliance_report_20260505_143022.xlsx")
    lines.append("```")
    lines.append("")
    lines.append("| Sheet | Contents |")
    lines.append("|-------|----------|")
    lines.append("| License Compliance | Entitled vs consumed with utilization % and status |")
    lines.append("| Monthly Users | Monthly protected object counts + consumed user calculation |")
    lines.append("| Monthly DPC (GB) | Monthly data transferred per workload type |")
    lines.append("| Monthly DPC (Bytes) | Raw byte values for precision calculations |")
    lines.append("| Per-Org Consumption | Current consumption broken down per M365 organization |")
    lines.append("| 10-Day Trend | Daily physical storage values for the last 10 days |")
    lines.append("")
    lines.append("> \u26a0\ufe0f Output files contain M365 organization names and data volumes. Treat as **Internal / Confidential** per your organization's data classification policy.")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Scheduling")
    lines.append("")
    lines.append("Run monthly to build a historical compliance archive.")
    lines.append("")
    lines.append("**Linux (cron):**")
    lines.append("")
    lines.append("```bash")
    lines.append("0 6 1 * * cd /path/to/project && ./venv/bin/python final_m365_report.py >> output/cron.log 2>&1")
    lines.append("```")
    lines.append("")
    lines.append("**Windows:** Use Task Scheduler \u2014 see the full operations guide at `output/RSC_M365_License_Compliance_Report_Guide.docx`")
    lines.append("")
    lines.append("**macOS:** Use launchd \u2014 see the full operations guide")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Project Structure")
    lines.append("")
    lines.append("```")
    lines.append("rsc-m365-storage-report/")
    lines.append("\u251c\u2500\u2500 final_m365_report.py    # Main report script")
    lines.append("\u251c\u2500\u2500 run_report.py           # Cross-platform entry-point wrapper")
    lines.append("\u251c\u2500\u2500 generate_docs.py        # Word doc + README generator")
    lines.append("\u251c\u2500\u2500 package_project.py      # Distribution packager")
    lines.append("\u251c\u2500\u2500 deploy.sh               # macOS/Linux deployment")
    lines.append("\u251c\u2500\u2500 deploy_linux.sh         # Linux deployment (distro-aware)")
    lines.append("\u251c\u2500\u2500 deploy.bat              # Windows deployment")
    lines.append("\u251c\u2500\u2500 requirements.txt        # Pinned Python dependencies")
    lines.append("\u251c\u2500\u2500 .env.example            # Configuration template (safe to commit)")
    lines.append("\u251c\u2500\u2500 .env                    # Your credentials \u2014 DO NOT COMMIT")
    lines.append("\u251c\u2500\u2500 SECURITY.md             # Vulnerability reporting policy")
    lines.append("\u251c\u2500\u2500 src/")
    lines.append("\u2502   \u251c\u2500\u2500 __init__.py")
    lines.append("\u2502   \u251c\u2500\u2500 auth.py             # RSC OAuth2 authentication")
    lines.append("\u2502   \u2514\u2500\u2500 graphql_client.py   # GraphQL query client")
    lines.append("\u2514\u2500\u2500 output/                 # Generated reports (excluded from git)")
    lines.append("```")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Security")
    lines.append("")
    lines.append("v2.1 was reviewed against **OWASP Top 10 (2021)**, **NIST CSF 2.0**, **CIS Controls v8**, and **MITRE ATT&CK for Enterprise**. The following hardening measures are in place.")
    lines.append("")
    lines.append("### Credential Protection")
    lines.append("")
    lines.append("- Credentials are loaded exclusively from `.env` \u2014 never hard-coded")
    lines.append("- Tokens are kept in memory only and never written to disk or logs")
    lines.append("- Exception messages are sanitized \u2014 `CLIENT_ID`, `CLIENT_SECRET`, and token values never appear in tracebacks")
    lines.append("- Use a **dedicated read-only service account**. Rotate the secret periodically to limit blast radius.")
    lines.append("")
    lines.append("### TLS Verification")
    lines.append("")
    lines.append("- All RSC API calls verify TLS against system CAs. This cannot be disabled.")
    lines.append("- `RSC_URL` is validated at startup against `https://*.my.rubrik.com` \u2014 plaintext URLs and unexpected hosts are rejected before any credentials are transmitted.")
    lines.append("- Corporate CA bundles can be specified via `RSC_CA_BUNDLE` without weakening verification.")
    lines.append("")
    lines.append("### Output File Security")
    lines.append("")
    lines.append("- `output/` is excluded from version control via `.gitignore`")
    lines.append("- Output files contain no credentials or tokens \u2014 only aggregated licensing and storage metrics")
    lines.append("")
    lines.append("### Dependency Auditing")
    lines.append("")
    lines.append("Dependencies are pinned in `requirements.txt` including `urllib3` and `certifi` for TLS-layer patch tracking. Keep them current:")
    lines.append("")
    lines.append("```bash")
    lines.append("pip install --upgrade -r requirements.txt")
    lines.append("```")
    lines.append("")
    lines.append("### Security Review Summary (v2.1)")
    lines.append("")
    lines.append("| # | Framework | Finding | Severity | Status |")
    lines.append("|---|-----------|---------|----------|--------|")
    lines.append("| 1 | OWASP A02 | `requests.post()` lacked explicit `verify=` \u2014 silent TLS bypass possible | High | \u2705 Fixed |")
    lines.append("| 2 | OWASP A07 | `os.getenv()` returned `None` silently \u2014 unauthenticated API calls possible | High | \u2705 Fixed |")
    lines.append("| 3 | OWASP A05 | `RSC_URL` accepted any scheme/host \u2014 credentials could be sent to wrong host | High | \u2705 Fixed |")
    lines.append("| 4 | OWASP A09 | No `raise_for_status()` \u2014 HTTP 401/403 silently returned `None` | High | \u2705 Fixed |")
    lines.append("| 5 | NIST PR.DS-1 | Exception tracebacks could expose credential fragments | Medium | \u2705 Fixed |")
    lines.append("| 6 | OWASP A03 | Nested `data[\"key\"]` access crashed on unexpected API response shapes | Medium | \u2705 Fixed |")
    lines.append("| 7 | NIST PR.DS-5 | `output/*.docx`, `*.pdf`, `*.log` not excluded from `.gitignore` | Medium | \u2705 Fixed |")
    lines.append("| 8 | CIS Control 4 | `urllib3` and `certifi` absent from `requirements.txt`; versions stale | Medium | \u2705 Fixed |")
    lines.append("| 9 | NIST PR.DS-2 | `os.makedirs` missing `exist_ok=True` \u2014 crash if `output/` absent | Low | \u2705 Fixed |")
    lines.append("| 10 | CIS Control 16 | `os.path.join` not used \u2014 forward-slash path non-portable on Windows | Low | \u2705 Fixed |")
    lines.append("")
    lines.append("See [SECURITY.md](SECURITY.md) for the vulnerability reporting policy.")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Troubleshooting")
    lines.append("")
    lines.append("| Issue | Solution |")
    lines.append("|-------|----------|")
    lines.append("| `RSC_URL is not set` | Copy `.env.example` to `.env` and set all three required variables |")
    lines.append("| `RSC_URL must use HTTPS` | Verify the URL in `.env` starts with `https://` |")
    lines.append("| `RSC_URL does not match expected pattern` | URL must be exactly `https://<org>.my.rubrik.com` |")
    lines.append("| TLS certificate verification failed | Set `RSC_CA_BUNDLE=/path/to/ca-bundle.crt` for corporate CAs |")
    lines.append("| `Authentication failed (HTTP 401)` | Verify `RSC_CLIENT_ID` and `RSC_CLIENT_SECRET` in `.env` |")
    lines.append("| `Authentication failed (HTTP 403)` | Service account needs read access to M365 objects and Reports |")
    lines.append("| Connection refused / Timeout | Check `RSC_URL`, confirm VPN/network, verify port 443 is open outbound |")
    lines.append("| Storage shows 0 | Normal for per-object M365 \u2014 use DPC / `transferredBytes` metrics instead |")
    lines.append("| No monthly data | Verify M365 backups are actively running in RSC |")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Documentation")
    lines.append("")
    lines.append("Full deployment and operations guide: `output/RSC_M365_License_Compliance_Report_Guide.docx`")
    lines.append("")
    lines.append("Generate with:")
    lines.append("")
    lines.append("```bash")
    lines.append("python generate_docs.py")
    lines.append("```")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Legal & Disclaimer")
    lines.append("")
    lines.append("This project is an **independent, open-source tool** and is **not affiliated with, authorized, maintained, sponsored, or endorsed by Rubrik, Inc.** in any way. All product and company names are the registered trademarks of their respective owners. The use of any trade name or trademark is for identification and reference purposes only.")
    lines.append("")
    lines.append("This software is provided **\"as-is,\" without warranty of any kind**. Use of this tool is at your own risk. The authors are not responsible for any data loss, API rate-limit overages, account suspensions, or security incidents resulting from the use of this software.")
    lines.append("")
    lines.append("You must have a valid API key and an active subscription or license for Rubrik Security Cloud (RSC). This software does not bypass any licensing checks or provide unauthorised access to Rubrik features.")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## License")
    lines.append("")
    lines.append("[Apache 2.0](LICENSE)")
    lines.append("")

    readme_path = "README.md"
    with open(readme_path, "w") as f:
        f.write("\n".join(lines))
    return readme_path


def main():
    print("=" * 60)
    print("  Generating Documentation & README")
    print("=" * 60 + "\n")

    print("Generating Word document...")
    doc_path = generate_word_doc()
    print(f"  Done: {doc_path}\n")

    print("Generating README.md...")
    readme_path = generate_readme()
    print(f"  Done: {readme_path}\n")

    print("=" * 60)
    print("  All documentation generated!")
    print(f"  Word:   {doc_path}")
    print(f"  README: {readme_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()